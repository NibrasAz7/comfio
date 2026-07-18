"""DOCX report export for IEQ results.

Provides :func:`ieq_to_docx` to generate a Word document with
KPIs, domain breakdown, and diagnostic insights. Requires ``python-docx``.
"""

from __future__ import annotations

import io

import numpy as np

from comfio.integration.global_ieq import GlobalIEQResult
from comfio.llm.interpreters import ieq_to_markdown
from comfio.performance.contracts import ComplianceReport

_IMPORT_ERROR_MSG = (
    "ieq_to_docx requires python-docx. Install it with: pip install python-docx\n"
    "Or install comfio with reports support: pip install comfio[gui-reports]"
)


def ieq_to_docx(
    ieq_result: GlobalIEQResult,
    compliance_report: ComplianceReport | None = None,
    zone_id: str | None = None,
) -> bytes:
    """Generate a Word document report from IEQ results.

    Parameters
    ----------
    ieq_result : GlobalIEQResult
        The Global IEQ Index calculation result.
    compliance_report : ComplianceReport or None
        Optional compliance report for compliance context.
    zone_id : str or None
        Optional zone identifier for the report title.

    Returns
    -------
    bytes
        DOCX file content.

    Raises
    ------
    ImportError
        If python-docx is not installed.

    Examples
    --------
    >>> import numpy as np
    >>> from comfio import evaluate_thermal, calculate_global_ieq
    >>> from comfio.reports import ieq_to_docx
    >>> thermal = evaluate_thermal(
    ...     tdb=np.array([24.0, 25.0]),
    ...     tr=np.array([24.0, 25.0]),
    ...     vr=np.array([0.1, 0.1]),
    ...     rh=np.array([50.0, 50.0]),
    ...     met=1.2, clo=0.5,
    ... )
    >>> ieq = calculate_global_ieq(thermal=thermal)
    >>> docx_bytes = ieq_to_docx(ieq)  # doctest: +SKIP
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(_IMPORT_ERROR_MSG) from exc

    doc = Document()

    # Title
    zone_label = f"Zone {zone_id}" if zone_id else "Building"
    doc.add_heading(f"IEQ Report: {zone_label}", level=0)

    # KPI summary
    doc.add_heading("KPI Summary", level=1)

    index_avg = float(np.mean(ieq_result.index))
    index_min = float(np.min(ieq_result.index))
    index_max = float(np.max(ieq_result.index))
    index_std = float(np.std(ieq_result.index))

    kpi_table = doc.add_table(rows=7, cols=2)
    kpi_table.style = "Light Grid Accent 1"
    kpi_rows = [
        ("Metric", "Value"),
        ("Global IEQ Index (avg)", f"{index_avg:.1f} / 100"),
        ("IEQ Index (min)", f"{index_min:.1f}"),
        ("IEQ Index (max)", f"{index_max:.1f}"),
        ("IEQ Index (std)", f"{index_std:.1f}"),
        ("Timestamps Evaluated", str(ieq_result.n_timestamps)),
        ("Domains Included", ", ".join(ieq_result.domains)),
    ]
    for row_idx, (label, value) in enumerate(kpi_rows):
        kpi_table.rows[row_idx].cells[0].text = label
        kpi_table.rows[row_idx].cells[1].text = value

    if compliance_report is not None:
        doc.add_paragraph(
            f"Compliance Rate: {compliance_report.compliance_rate_pct:.1f}% "
            f"(threshold: {compliance_report.threshold:.0f})"
        )

    # Domain breakdown
    doc.add_heading("Domain Breakdown", level=1)

    domain_avgs = {d: float(np.mean(s)) for d, s in ieq_result.domain_scores.items()}

    dom_table = doc.add_table(rows=len(domain_avgs) + 1, cols=4)
    dom_table.style = "Light Grid Accent 1"
    dom_table.rows[0].cells[0].text = "Domain"
    dom_table.rows[0].cells[1].text = "Avg Score"
    dom_table.rows[0].cells[2].text = "Weight"
    dom_table.rows[0].cells[3].text = "Status"

    for row_idx, (domain, avg) in enumerate(domain_avgs.items(), start=1):
        weight = ieq_result.weights_used.get(domain, 0.0)
        status = "OK" if avg >= 70 else "WARNING"
        dom_table.rows[row_idx].cells[0].text = domain.upper()
        dom_table.rows[row_idx].cells[1].text = f"{avg:.1f}"
        dom_table.rows[row_idx].cells[2].text = f"{weight:.2f}"
        dom_table.rows[row_idx].cells[3].text = status

    # Diagnostic
    doc.add_heading("Diagnostic Insight", level=1)
    worst_domain = min(domain_avgs, key=lambda k: domain_avgs[k]) if domain_avgs else "N/A"
    worst_score = domain_avgs.get(worst_domain, 0.0)
    doc.add_paragraph(
        f"The primary limiting factor is the {worst_domain.upper()} domain "
        f"(score: {worst_score:.1f}/100)."
    )

    # Full markdown report
    doc.add_heading("Full Report", level=1)
    md = ieq_to_markdown(ieq_result, compliance_report=compliance_report, zone_id=zone_id)
    for line in md.split("\n"):
        line = line.strip()
        if not line:
            continue
        clean = line.replace("*", "").replace("#", "")
        doc.add_paragraph(clean)

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
